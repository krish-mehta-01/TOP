"""
build_docx.py -- regenerates "TOP1 updated .docx" directly from the same
rendered TOP_Research_Paper.html that build.py produces from the Markdown
source, so the Word copy always matches the PDF. SVG figures are swapped for
their rasterized PNG twins (figures/<name>.png, produced by
render_svgs_to_png.py) since Word cannot embed SVG.

Usage: python build.py && python render_svgs_to_png.py && python build_docx.py
"""

import os
import re

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
HTML_PATH = os.path.join(HERE, "TOP_Research_Paper.html")
FIG_DIR = os.path.join(HERE, "figures")
OUT_PATH = os.path.join(HERE, "TOP1 updated .docx")

BODY_FONT = "Times New Roman"
MONO_FONT = "Courier New"
MAX_IMG_WIDTH_IN = 6.3
MAX_IMG_HEIGHT_IN = 6.2


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11.5)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def add_runs_from_inline(paragraph, node, base_bold=False, base_italic=False):
    """Walk inline children of an HTML node (strong/em/code/text) and add
    matching runs to a docx paragraph, preserving bold/italic/monospace."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text == "":
                continue
            run = paragraph.add_run(text)
            run.bold = base_bold
            run.italic = base_italic
        elif isinstance(child, Tag):
            if child.name == "strong":
                add_runs_from_inline(paragraph, child, True, base_italic)
            elif child.name == "em":
                add_runs_from_inline(paragraph, child, base_bold, True)
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = MONO_FONT
                run.font.size = Pt(10)
                run.bold = base_bold
                run.italic = base_italic
            elif child.name == "br":
                paragraph.add_run().add_break()
            else:
                add_runs_from_inline(paragraph, child, base_bold, base_italic)


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(17)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pPr.makeelement(
            qn("w:bottom"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "2", qn("w:color"): "999999"},
        )
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11.5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    return p


def resolve_image_path(src):
    src = src.replace("/", os.sep)
    if src.lower().endswith(".svg"):
        png = src[:-4] + ".png"
        candidate = os.path.join(HERE, png)
        if os.path.exists(candidate):
            return candidate
    return os.path.join(HERE, src)


def add_image(doc, src, caption_text=None):
    path = resolve_image_path(src)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        p.add_run(f"[missing image: {src}]").italic = True
        return
    with Image.open(path) as im:
        w_px, h_px = im.size
    aspect = h_px / w_px
    width_in = MAX_IMG_WIDTH_IN
    if width_in * aspect > MAX_IMG_HEIGHT_IN:
        width_in = MAX_IMG_HEIGHT_IN / aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption_text)
        cap_run.italic = True
        cap_run.font.size = Pt(9.5)
        cap.paragraph_format.space_after = Pt(10)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcPr.append(shd)


def add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    n_cols = len(rows[0].find_all(["th", "td"]))
    doc_table = doc.add_table(rows=0, cols=n_cols)
    doc_table.style = "Table Grid"

    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        is_header = tr.find("th") is not None
        row = doc_table.add_row()
        for c_idx, cell in enumerate(cells):
            if c_idx >= n_cols:
                continue
            docx_cell = row.cells[c_idx]
            docx_cell.paragraphs[0].text = ""
            p = docx_cell.paragraphs[0]
            add_runs_from_inline(p, cell, base_bold=is_header)
            for run in p.runs:
                run.font.size = Pt(9.3)
            if is_header:
                set_cell_shading(docx_cell, "E8E8E8")
    return doc_table


def walk_body(doc, body):
    children = list(body.children)
    i = 0
    while i < len(children):
        node = children[i]
        i += 1
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue

        if node.name == "h1":
            add_heading(doc, node.get_text(), 1)
        elif node.name == "h2":
            add_heading(doc, node.get_text(), 2)
        elif node.name == "h3":
            add_heading(doc, node.get_text(), 3)
        elif node.name == "hr":
            continue
        elif node.name == "p":
            imgs = node.find_all("img")
            if imgs:
                for img in imgs:
                    add_image(doc, img.get("src", ""), img.get("alt", "") or None)
                continue
            text = node.get_text(strip=True)
            if not text:
                continue
            p = doc.add_paragraph()
            add_runs_from_inline(p, node)
        elif node.name in ("ul", "ol"):
            style = "List Bullet" if node.name == "ul" else "List Number"
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                add_runs_from_inline(p, li)
        elif node.name == "table":
            add_table(doc, node)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(8)
        elif node.name == "blockquote":
            p = doc.add_paragraph()
            add_runs_from_inline(p, node)
            p.paragraph_format.left_indent = Cm(0.6)
        elif node.name == "div":
            # gallery / grouping divs (none expected post-appendix-removal); recurse just in case
            walk_body(doc, node)
        else:
            text = node.get_text(strip=True)
            if text:
                p = doc.add_paragraph()
                add_runs_from_inline(p, node)


def main():
    with open(HTML_PATH, encoding="utf-8") as f:
        html = f.read()
    soup = BeautifulSoup(html, "html.parser")
    body = soup.find("body")

    doc = setup_document()
    walk_body(doc, body)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
